# Copyright (c) 2023 Arista Networks, Inc.
# Use of this source code is governed by the Apache License 2.0
# that can be found in the LICENSE file.
"""
Word report management
"""
from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt
from docx.table import Table

from anta.reporter.word_models import Color, Colors, DocxStyle, TextCover, TextInputs, TextSection
from anta.result_manager import ResultManager
from anta.result_manager.models import TestResult

logger = logging.getLogger(__name__)


WIDTHS = (
    Inches(0.5),
    Inches(1.17),
    Inches(2.50),
    Inches(1.25),
    Inches(0.75),
    Inches(1.50),
)

DEFAULT_REPORT_FORMAT = TextInputs(
    cover=TextCover(),
    content=[
        TextSection(index=1, header="Abstract"),
        TextSection(index=2, header="Test Description", insert="tests_description", break_after=True),
        TextSection(
            index=3,
            header="Summary per tests",
            description=["Ad minim est reprehenderit id adipisicing quis sunt minim id tempor dolore."],
            insert="summary_per_test",
            break_after=True,
        ),
        TextSection(
            index=4,
            header="Tests per device",
            description=["Ad minim est reprehenderit id adipisicing quis sunt minim id tempor dolore."],
            insert="tests_per_device",
        ),
    ],
)

DEFAULT_REPORT_STYLE = DocxStyle(
    colors=Colors(
        headers_table=Color(hex="#1F618D"),
        success=Color(hex="#A9DFBF"),
        failure=Color(hex="#F5B7B1"),
        error=Color(hex="#E74C3C"),
        skipped=Color(hex="#EB984E"),
    ),
    font_colors=Colors(headers_table=Color(hex="#FFFFFF")),
    font_size_default=9,
    font_name="Arial",
    table_style="Table Grid",
    shade_errors=True,
    shade_success=False,
    shade_failure=True,
)


class ReportWordDocx:
    """Class to manage docx reporting."""

    def __init__(self, filename: Path, anta_result_manager: ResultManager, docx_style: Optional[DocxStyle] = None) -> None:
        """Class constructor."""
        # Baseline
        self.anta_results = anta_result_manager
        self.docx_style = docx_style if docx_style is not None else DEFAULT_REPORT_STYLE
        self.filename = filename
        self.document = Document()

    # -------------------------------------------
    # Internal methods
    # -------------------------------------------

    def _set_header_generic(self, level: int, content: str, font_size: int) -> None:
        """
        Wrapper to insert a title with a custom level and font size.
        """
        heading = self.document.add_heading(level=level)
        run = heading.add_run(content)
        run.font.size = Pt(font_size)
        # run.font.color.rgb = RGBColor(204, 0, 0)

    def _shade_cell(self, cell_idx: int, table: Table, count: int, shade: str) -> None:
        """
        Shade a cell in word doc table

        Args:
            cell_idx (int): Column index for cell to shade
            table (obj): Word doc table object
            count (int): Row index for cell to shade
            shade (str): hexadecimal color representation
        """
        cell = table.cell(count, cell_idx)
        color = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shade}"/>')
        # pylint: disable-next=protected-access
        cell._tc.get_or_add_tcPr().append(color)

    def _create_table(self, headers: List[str]) -> Table:
        """Create a table with its headers."""
        # add table ------------------
        table = self.document.add_table(rows=1, cols=len(headers), style=self.docx_style.table_style)
        table.autofit = True

        for count, header in enumerate(headers):
            para = table.rows[0].cells[count].paragraphs[0]
            run = para.add_run(header)
            run.font.bold = True
            run.font.name = self.docx_style.font_name
            run.font.size = Pt(self.docx_style.font_size_default)
            run.font.color.rgb = self.docx_style.font_colors.headers_table.rgb()
            self._shade_cell(count, table, 0, self.docx_style.colors.headers_table.hex)

        return table

    def _clean_table(self, table: Table) -> None:
        """Force all entries to use correct font settings."""
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(self.docx_style.font_size_default)

    def _count_test_results(self, result: List[TestResult]) -> Dict[str, Any]:
        """Count result types in a list of TestResult."""
        # TODO: should be part of ResultManager to avoid duplicate code.
        return {
            "success": len([e for e in result if e.result == "success"]),
            "failure": len([e for e in result if e.result == "failure"]),
            "error": len([e for e in result if e.result == "error"]),
            "skipped": len([e for e in result if e.result == "skipped"]),
            "devices": [e.name for e in result if e.result in ["error", "failure"]],
        }

    # -------------------------------------------
    # Pubic methods
    # -------------------------------------------

    def save(self) -> None:
        """Close and Save document."""
        logger.info(f"Saving report under {self.filename.absolute()}")
        self.document.save(str(self.filename.absolute()))

    def add_header_l1(self, content: str) -> None:
        """Insert a level 1 header."""
        self._set_header_generic(level=1, content=content, font_size=24)
        # run.font.color.rgb = RGBColor(204, 0, 0)

    def add_header_l2(self, content: str) -> None:
        """Insert a level 2 header."""
        self._set_header_generic(level=2, content=content, font_size=18)
        # run.font.color.rgb = RGBColor(204, 0, 0)

    def add_header_l3(self, content: str) -> None:
        """Insert a level 3 header."""
        self._set_header_generic(level=3, content=content, font_size=14)

    def add_text(self, content: str) -> None:
        """Add paragraph to document."""
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
        run = para.add_run(content)
        run.font.size = Pt(self.docx_style.font_size_default)
        run.font.name = self.docx_style.font_name

    def generate_tests_description(self) -> None:
        """Generate a section with a table with test names and their descriptions."""
        testnames = self.anta_results.get_testcases()
        tests_description = []
        for testname in testnames:
            desc = [d["description"] for d in json.loads(self.anta_results.get_json_results()) if d["test"] == testname][0]
            tests_description.append({"test": f"{testname}", "description": f"{desc}"})

        headers = ["Test Name", "Description"]
        table = self._create_table(headers)
        for item in tests_description:
            cells = table.add_row().cells
            cells[0].text = item["test"]
            cells[1].text = item["description"]

        self._clean_table(table)

    def generate_summary_per_test(self) -> None:
        """Generate a section with a summary of result grouped by tests name."""
        headers = ["Test Case", "# of success", "# of skipped", "# of failure", "# of erros", "List of failed or error nodes"]
        tests_name = self.anta_results.get_testcases()

        table = self._create_table(headers=headers)
        for i, tname in enumerate(tests_name, start=1):
            result = self.anta_results.get_result_by_test(tname)
            result_numbers = self._count_test_results(result)
            row_cells = table.add_row().cells
            row_cells[0].text = tname
            row_cells[5].text = str(result_numbers["devices"])

            # Manage Cell for # of success
            run_success = row_cells[1].paragraphs[0].add_run(str(result_numbers["success"]))
            if result_numbers["success"] > 0:
                run_success.font.color.rgb = self.docx_style.colors.success.rgb()

            run_success = row_cells[2].paragraphs[0].add_run(str(result_numbers["skipped"]))
            if result_numbers["skipped"] > 0:
                if self.docx_style.shade_success:
                    self._shade_cell(3, table, i, self.docx_style.colors.skipped.hex)
                else:
                    run_success.font.color.rgb = self.docx_style.colors.skipped.rgb()

            run_success = row_cells[3].paragraphs[0].add_run(str(result_numbers["failure"]))
            if result_numbers["failure"] > 0:
                if self.docx_style.shade_failure:
                    self._shade_cell(3, table, i, self.docx_style.colors.failure.hex)
                else:
                    run_success.font.color.rgb = self.docx_style.colors.failure.rgb()

            run_success = row_cells[4].paragraphs[0].add_run(str(result_numbers["error"]))
            if result_numbers["error"] > 0:
                if self.docx_style.shade_errors:
                    self._shade_cell(3, table, i, self.docx_style.colors.error.hex)
                else:
                    run_success.font.color.rgb = self.docx_style.colors.error.rgb()

    def generate_tests_report(self, device_name: str, index: int, index_sub: int = 1) -> None:
        """Populate table with summary of tests run by ANTA."""
        test_data = self.anta_results.get_result_by_host(device_name)
        headers = [
            "Device name",
            "Case Name",
            "Pass/Fail",
            "Observation",
        ]

        self.add_header_l3(f"{index}.{index_sub} Tests for device {device_name}")
        table = self._create_table(headers)

        for count, test_case_entry in enumerate(test_data, start=1):
            row_cells = table.add_row().cells
            # Add entries with no formatting
            row_cells[0].text = test_case_entry.name
            row_cells[1].text = test_case_entry.test
            row_cells[3].text = test_case_entry.messages

            # Add entries using formatting
            run = row_cells[2].paragraphs[0].add_run(test_case_entry.result)
            if test_case_entry.result == "failure":
                if self.docx_style.shade_failure:
                    self._shade_cell(2, table, count, self.docx_style.colors.failure.hex)
                else:
                    run.font.color.rgb = self.docx_style.colors.failure.rgb()
            elif test_case_entry.result == "success":
                if self.docx_style.shade_success:
                    self._shade_cell(2, table, count, self.docx_style.colors.success.hex)
                else:
                    run.font.color.rgb = self.docx_style.colors.success.rgb()
            run.bold = True
            run.font.name = self.docx_style.font_name

        self._clean_table(table)

    def report_template(self, text_report: Optional[TextInputs] = None) -> None:
        """Generic function to render WORD report based on TextInputs."""
        if text_report is None:
            text_report = DEFAULT_REPORT_FORMAT
        # Run Header
        self.add_text(content="\n\n\n\n\n\n\n\n\n\n\n\n\n\n")
        self.add_header_l1(text_report.cover.title)
        for subtitle in text_report.cover.subtitles:
            self.add_text(content=subtitle)
        self.document.add_page_break()

        for section in text_report.content:
            self.add_header_l2(content=f"{section.index}. {section.header}")
            # Add description blocks
            for descr in section.description:
                self.add_text(descr)

            # Add insertion
            if section.insert == "tests_description":
                self.generate_tests_description()
            elif section.insert == "summary_per_test":
                self.generate_summary_per_test()
            elif section.insert == "tests_per_device":
                for sub_index, device in enumerate(self.anta_results.get_hosts(), start=1):
                    self.generate_tests_report(device_name=device, index=section.index, index_sub=sub_index)

            # Page Break once section is added
            if section.break_after:
                self.document.add_page_break()

        self.save()
